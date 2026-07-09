import base64
from dataclasses import dataclass
import io
from pathlib import Path
from typing import Any
from urllib.parse import quote, urljoin, urlparse, urlunparse
import warnings

from bs4 import BeautifulSoup
import markdown
import requests

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# 代码块 / 行内代码样式常量
_CODE_FONT = "Consolas"
_CODE_BLOCK_BG = "F5F5F5"
_CODE_BLOCK_ACCENT = "2F75B5"
_CODE_INLINE_BG = "EEEEEE"

# 表格样式常量
_TABLE_HEADER_BG = "2F75B5"
_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
_TABLE_BORDER = "BFBFBF"
_TABLE_BODY_ALT_BG = "F2F6FA"

# 标题默认样式（h1–h6 对应索引 0–5）
_HEADING_FONT_SIZES_PT = (18, 16, 14, 13, 12, 11)
_HEADING_SPACE_BEFORE_PT = (24, 18, 14, 12, 10, 8)
_HEADING_SPACE_AFTER_PT = (18, 14, 12, 10, 8, 6)

# 正文样式：宋体小四（12pt）；西文 Times New Roman
_BODY_FONT = "宋体"
_BODY_FONT_SIZE_PT = 12
_LATIN_FONT = "Times New Roman"
_HYPERLINK_COLOR = "0563C1"


@dataclass
class HeadingStyleConfig:
    """标题编号与样式配置。"""

    numbering_enabled: bool = True
    use_word_numbering: bool = True
    max_numbering_level: int = 6
    number_separator: str = "."
    number_suffix: str = " "
    trailing_dot_on_top_level: bool = True

    font_name: str = "黑体"
    font_sizes_pt: tuple[float, ...] = _HEADING_FONT_SIZES_PT
    space_before_pt: tuple[float, ...] = _HEADING_SPACE_BEFORE_PT
    space_after_pt: tuple[float, ...] = _HEADING_SPACE_AFTER_PT

    def level_value(self, level: int, values: tuple[float, ...], default: float) -> float:
        idx = min(max(level, 1), 6) - 1
        return values[idx] if idx < len(values) else default


class _HeadingNumbering:
    def __init__(self, config: HeadingStyleConfig):
        self.config = config
        self._counters = [0] * 6

    def reset(self):
        self._counters = [0] * 6

    def next(self, level: int) -> str:
        if not self.config.numbering_enabled:
            return ""
        level = min(max(level, 1), 6)
        if level > self.config.max_numbering_level:
            return ""

        idx = level - 1
        self._counters[idx] += 1
        for i in range(idx + 1, 6):
            self._counters[i] = 0

        parts = [str(self._counters[i]) for i in range(idx + 1)]
        number = self.config.number_separator.join(parts)
        if self.config.trailing_dot_on_top_level and level == 1:
            number += "."
        return number + self.config.number_suffix


def _set_element_font_black(r_pr):
    """移除主题色，强制黑色字体。"""
    color_elem = r_pr.find(qn("w:color"))
    if color_elem is not None:
        r_pr.remove(color_elem)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)


def _set_element_fonts(
    r_pr,
    *,
    east_asia: str,
    latin: str = _LATIN_FONT,
    bold: bool | None = None,
):
    """西文 latin + 东亚 east_asia 分字体（Word 按字符自动选用）。"""
    if bold is True:
        b_elem = r_pr.find(qn("w:b"))
        if b_elem is None:
            r_pr.append(OxmlElement("w:b"))
    elif bold is False:
        b_elem = r_pr.find(qn("w:b"))
        if b_elem is not None:
            r_pr.remove(b_elem)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


'''
来自Dify的Markdown to DOCX Converter的工具：Repository: https://github.com/stvlynn/DOC-Dify-Plugin.

核心依赖："python-docx>=1.2.0","markdown>=3.10.1",beautifulsoup4
'''


class DocTool():
    def __init__(self, heading_config: HeadingStyleConfig | None = None):
        self.heading_config = heading_config or HeadingStyleConfig()
        self._heading_numbering = _HeadingNumbering(self.heading_config)
        self._heading_num_id: int | None = None

    def create_text_message(self, content: str):
        """创建文本消息"""
        return {"type": "text", "content": content}
    
    def create_blob_message(self, blob: bytes, meta: dict):
        """创建二进制数据消息"""
        return {"type": "blob", "data": blob, "meta": meta}
    
    def invoke(self, tool_parameters: dict[str, Any]):
        # Get markdown content from parameters
        md_content = tool_parameters.get("markdown_content", "")
        title = tool_parameters.get("title", "Document")
        base_path = tool_parameters.get("base_path")
        heading_config = tool_parameters.get("heading_config")
        
        if not md_content:
            yield self.create_text_message("No markdown content provided.")
            return
        
        try:
            # Create a new Word document
            doc = self._convert_markdown_to_docx(
                md_content,
                title,
                base_path=base_path,
                heading_config=heading_config,
            )
            
            # Save document to a bytes buffer
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            # Get the byte data
            file_bytes = docx_bytes.getvalue()
            
            # Create a filename
            filename = f"{title.replace(' ', '_')}.docx"
            
            # Return success message
            yield self.create_text_message(f"Document '{title}' generated successfully")
            
            # Return the document data as a blob
            yield self.create_blob_message(
                blob=file_bytes, 
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": filename
                }
            )
        except Exception as e:
            yield self.create_text_message(f"Error converting markdown to DOCX: {str(e)}")
    
    def _convert_markdown_to_docx(
        self,
        md_content: str,
        title: str,
        base_path: str | Path | None = None,
        heading_config: HeadingStyleConfig | None = None,
    ) -> Document:
        if heading_config is not None:
            self.heading_config = heading_config
            self._heading_numbering = _HeadingNumbering(heading_config)
        else:
            self._heading_numbering.reset()

        self._heading_num_id = None
        doc = Document()
        if (
            self.heading_config.numbering_enabled
            and self.heading_config.use_word_numbering
        ):
            self._heading_num_id = self._setup_heading_numbering(doc)
        self._configure_heading_styles(doc)
        self._configure_body_styles(doc)

        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.name = _LATIN_FONT
        title_run.font.size = Pt(18)
        _set_element_fonts(
            title_run._element.get_or_add_rPr(),
            east_asia=self.heading_config.font_name,
            bold=True,
        )
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Convert markdown to HTML with extensions
        html_content = markdown.markdown(
            md_content, 
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.codehilite',
                'markdown.extensions.nl2br',
                'markdown.extensions.sane_lists'
            ]
        )
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Process HTML elements
        self._process_html_elements(doc, soup, base_path=base_path)
        
        return doc
    
    def _process_html_elements(self, doc, soup, base_path: str | Path | None = None):
        # Process all elements
        for element in soup.children:
            if element.name is None:
                continue
                
            elif element.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(element.name[1])
                self._add_heading(doc, element.get_text(), level)
                
            elif element.name == 'p':
                p = doc.add_paragraph()
                self._add_run_with_formatting(p, element, base_path=base_path)
                
            elif element.name == 'ul':
                self._add_list(doc, element, is_numbered=False, base_path=base_path)
                
            elif element.name == 'ol':
                self._add_list(doc, element, is_numbered=True, base_path=base_path)
                
            elif element.name == 'pre':
                code, lang = self._extract_code_from_pre(element)
                self._add_code_block(doc, code, lang)
                
            elif element.name == 'table':
                self._add_html_table(doc, element)

            elif element.name == 'hr':
                continue  # Markdown --- 分隔线不输出到 Word
                
            elif element.name == 'img':
                self._add_image(doc, element, base_path=base_path)
                
            elif hasattr(element, 'children'):
                # Recursively process child elements
                self._process_html_elements(doc, element, base_path=base_path)

    def _build_heading_lvl_text(self, ilvl: int) -> str:
        config = self.heading_config
        parts = [f"%{i + 1}" for i in range(ilvl + 1)]
        text = config.number_separator.join(parts)
        if config.trailing_dot_on_top_level and ilvl == 0:
            text += "."
        return text + config.number_suffix

    def _apply_heading_num_pr(self, p_pr, num_id: int, ilvl: int):
        existing = p_pr.find(qn("w:numPr"))
        if existing is not None:
            p_pr.remove(existing)
        num_pr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_pr.append(ilvl_el)
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_el)
        p_pr.append(num_pr)

    def _remove_heading_num_pr(self, p_pr):
        existing = p_pr.find(qn("w:numPr"))
        if existing is not None:
            p_pr.remove(existing)

    def _setup_heading_numbering(self, doc: Document) -> int:
        """在 numbering.xml 中创建标题多级列表，返回 numId。"""
        config = self.heading_config
        numbering = doc.part.numbering_part.element

        abstract_ids = [
            int(el.get(qn("w:abstractNumId")))
            for el in numbering.findall(qn("w:abstractNum"))
        ]
        abstract_num_id = max(abstract_ids, default=-1) + 1

        num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
        num_id = max(num_ids, default=0) + 1

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract_num.append(multi)

        max_level = min(config.max_numbering_level, 6)
        for ilvl in range(max_level):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(ilvl))

            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)

            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "decimal")
            lvl.append(num_fmt)

            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), self._build_heading_lvl_text(ilvl))
            lvl.append(lvl_text)

            p_style = OxmlElement("w:pStyle")
            p_style.set(qn("w:val"), f"Heading{ilvl + 1}")
            lvl.append(p_style)

            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            lvl.append(lvl_jc)

            abstract_num.append(lvl)

        numbering.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id_el = OxmlElement("w:abstractNumId")
        abstract_num_id_el.set(qn("w:val"), str(abstract_num_id))
        num.append(abstract_num_id_el)
        numbering.append(num)

        return num_id

    def _configure_heading_styles(self, doc: Document):
        """配置 Word 内置 Heading 1–6 样式（黑体、黑色、字号与间距）。"""
        config = self.heading_config
        for level in range(1, 7):
            style = doc.styles[f"Heading {level}"]
            idx = level - 1

            font = style.font
            font.name = _LATIN_FONT
            font.size = Pt(config.level_value(level, config.font_sizes_pt, 11))
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)

            r_pr = style.element.get_or_add_rPr()
            _set_element_fonts(r_pr, east_asia=config.font_name, bold=True)
            _set_element_font_black(r_pr)

            sz = r_pr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                r_pr.append(sz)
            half_points = str(int(config.level_value(level, config.font_sizes_pt, 11) * 2))
            sz.set(qn("w:val"), half_points)

            pf = style.paragraph_format
            pf.space_before = Pt(config.level_value(level, config.space_before_pt, 8))
            pf.space_after = Pt(config.level_value(level, config.space_after_pt, 6))
            pf.line_spacing = 1.15
            pf.keep_with_next = True

            # 确保大纲级别与标题层级一致（Heading 1 → outlineLvl 0）
            p_pr = style.element.get_or_add_pPr()
            outline = p_pr.find(qn("w:outlineLvl"))
            if outline is None:
                outline = OxmlElement("w:outlineLvl")
                p_pr.append(outline)
            outline.set(qn("w:val"), str(idx))

            if (
                self._heading_num_id is not None
                and level <= config.max_numbering_level
            ):
                self._apply_heading_num_pr(p_pr, self._heading_num_id, idx)
            else:
                self._remove_heading_num_pr(p_pr)

    def _configure_body_styles(self, doc: Document):
        """配置正文样式：宋体小四。"""
        for style_name in ("Normal", "List Paragraph", "List Bullet", "List Number"):
            try:
                style = doc.styles[style_name]
            except KeyError:
                continue

            font = style.font
            font.name = _LATIN_FONT
            font.size = Pt(_BODY_FONT_SIZE_PT)
            font.bold = False
            font.color.rgb = RGBColor(0, 0, 0)

            r_pr = style.element.get_or_add_rPr()
            _set_element_fonts(r_pr, east_asia=_BODY_FONT, bold=False)
            _set_element_font_black(r_pr)

            sz = r_pr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                r_pr.append(sz)
            sz.set(qn("w:val"), str(int(_BODY_FONT_SIZE_PT * 2)))

    def _apply_body_run_font(self, run):
        run.font.name = _LATIN_FONT
        run.font.size = Pt(_BODY_FONT_SIZE_PT)
        run.font.color.rgb = RGBColor(0, 0, 0)
        r_pr = run._element.get_or_add_rPr()
        _set_element_fonts(r_pr, east_asia=_BODY_FONT, bold=False)
        _set_element_font_black(r_pr)

    def _apply_heading_run_font(self, run, east_asia_font: str, *, size_pt: float):
        run.font.name = _LATIN_FONT
        run.bold = True
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)
        r_pr = run._element.get_or_add_rPr()
        _set_element_fonts(r_pr, east_asia=east_asia_font, bold=True)
        _set_element_font_black(r_pr)

    def _add_body_run(self, paragraph, text: str, *, bold: bool = False, italic: bool = False):
        run = paragraph.add_run(text)
        self._apply_body_run_font(run)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return run

    def _resolve_link_url(self, href: str, base_path: str | Path | None = None) -> str:
        href = (href or "").strip()
        if not href:
            return href
        parsed = urlparse(href)
        if parsed.scheme in ("http", "https", "mailto", "ftp"):
            return href
        if base_path:
            base_text = str(base_path)
            if base_text.startswith(("http://", "https://")):
                return urljoin(base_text.rstrip("/") + "/", href)
            return str(Path(base_path) / href)
        return href

    def _add_hyperlink(self, paragraph, text: str, url: str):
        text = text or ""
        url = (url or "").strip()
        if not text:
            return
        if not url:
            self._add_body_run(paragraph, text)
            return

        r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        _set_element_fonts(r_pr, east_asia=_BODY_FONT, bold=False)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), _HYPERLINK_COLOR)
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(_BODY_FONT_SIZE_PT * 2)))
        r_pr.append(sz)
        run.append(r_pr)

        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = text
        run.append(text_el)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _append_inline_node(self, paragraph, node, base_path: str | Path | None = None):
        if node.name is None:
            self._add_body_run(paragraph, node.string or "")
            return

        if node.name in ("strong", "b"):
            self._add_body_run(paragraph, node.get_text(), bold=True)
        elif node.name in ("em", "i"):
            self._add_body_run(paragraph, node.get_text(), italic=True)
        elif node.name == "code":
            self._add_inline_code_run(paragraph, node.get_text())
        elif node.name == "a":
            href = self._resolve_link_url(node.get("href", ""), base_path=base_path)
            self._add_hyperlink(paragraph, node.get_text(), href)
        elif node.name == "img":
            src = node.get("src")
            alt = node.get("alt", "")
            if src:
                try:
                    image_stream = self._load_image_stream(src, base_path=base_path)
                    paragraph.add_run().add_picture(image_stream)
                    if alt:
                        self._add_body_run(paragraph, f" ({alt})", italic=True)
                except Exception:
                    self._add_body_run(paragraph, f"[图片加载失败: {alt or src}]", italic=True)
        elif node.name == "br":
            paragraph.add_run().add_break()
        elif hasattr(node, "children"):
            for child in node.children:
                self._append_inline_node(paragraph, child, base_path=base_path)

    def _add_heading(self, doc: Document, text: str, level: int):
        config = self.heading_config
        level = min(max(level, 1), 6)
        text = text.strip()
        if not text:
            return

        use_word_num = (
            config.numbering_enabled
            and config.use_word_numbering
            and self._heading_num_id is not None
            and level <= config.max_numbering_level
        )
        if use_word_num:
            display_text = text
        elif config.numbering_enabled and level <= config.max_numbering_level:
            number = self._heading_numbering.next(level)
            display_text = f"{number}{text}" if number else text
        else:
            display_text = text

        paragraph = doc.add_heading(display_text, level=level)
        if use_word_num:
            self._apply_heading_num_pr(
                paragraph._element.get_or_add_pPr(),
                self._heading_num_id,
                level - 1,
            )
        for run in paragraph.runs:
            self._apply_heading_run_font(
                run,
                config.font_name,
                size_pt=config.level_value(level, config.font_sizes_pt, 11),
            )
    
    def _add_run_with_formatting(self, paragraph, element, base_path: str | Path | None = None):
        for child in element.children:
            self._append_inline_node(paragraph, child, base_path=base_path)
    
    def _add_list(self, doc, list_element, is_numbered=False, base_path: str | Path | None = None):
        # Process a list (ul or ol)
        for item in list_element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number' if is_numbered else 'List Bullet')
            self._add_run_with_formatting(p, item, base_path=base_path)
            
            # Process nested lists
            nested_ul = item.find('ul')
            nested_ol = item.find('ol')
            
            if nested_ul:
                self._add_list(doc, nested_ul, is_numbered=False, base_path=base_path)
            if nested_ol:
                self._add_list(doc, nested_ol, is_numbered=True, base_path=base_path)

    def _decode_base64_image(self, data_url: str) -> io.BytesIO:
        """解析 data:image/...;base64,... 并返回图片流。"""
        if "," not in data_url:
            raise ValueError("Invalid data URL format")
        header, encoded_data = data_url.split(",", 1)
        if ";base64" not in header.lower():
            raise ValueError("Unsupported data URL: only base64 image is supported")
        try:
            return io.BytesIO(base64.b64decode(encoded_data))
        except Exception as e:
            raise ValueError(f"Invalid base64 image data: {e}") from e

    def _load_local_image(self, image_path: str, base_path: str | Path | None = None) -> io.BytesIO:
        """从本地路径读取图片并返回字节流。"""
        candidate_path = Path(image_path)
        if not candidate_path.is_absolute():
            if base_path is None:
                raise FileNotFoundError(
                    f"Relative image path requires base_path, got: {image_path}"
                )
            candidate_path = Path(base_path) / candidate_path

        if not candidate_path.exists():
            raise FileNotFoundError(f"Image file not found: {candidate_path}")
        if not candidate_path.is_file():
            raise ValueError(f"Image path is not a file: {candidate_path}")
        return io.BytesIO(candidate_path.read_bytes())

    def _load_image_stream(self, src: str, base_path: str | Path | None = None) -> io.BytesIO:
        """统一处理图片来源（base64、URL、本地路径）并返回字节流。"""
        src = src.strip()
        if not src:
            raise ValueError("Empty image source")

        if src.lower().startswith("data:image/"):
            return self._decode_base64_image(src)

        parsed = urlparse(src)
        if parsed.scheme in ("http", "https"):
            return self._download_image(src)
        if parsed.scheme == "file":
            return self._load_local_image(parsed.path, base_path=None)

        if base_path:
            base_text = str(base_path)
            if base_text.startswith(("http://", "https://")):
                resolved_url = urljoin(base_text.rstrip("/") + "/", src)
                return self._download_image(resolved_url)

        return self._load_local_image(src, base_path=base_path)
    
    def _extract_code_from_pre(self, pre_element) -> tuple[str, str]:
        """从 <pre>（及可选的 <code>）提取纯文本与语言标识。"""
        code_el = pre_element.find("code")
        target = code_el or pre_element
        lang = ""
        for cls in target.get("class") or []:
            if cls.startswith("language-"):
                lang = cls[9:]
                break
            if cls not in ("highlight", "codehilite", "sourceCode"):
                lang = cls
                break
        return target.get_text(), lang

    def _set_run_mono_font(self, run, *, size_pt: float = 9, inline: bool = False):
        """设置等宽字体（含东亚字体回退）。"""
        run.font.name = _CODE_FONT
        run.font.size = Pt(8 if inline else size_pt)
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), _CODE_FONT)
        r_fonts.set(qn("w:hAnsi"), _CODE_FONT)
        r_fonts.set(qn("w:cs"), _CODE_FONT)
        r_fonts.set(qn("w:eastAsia"), _CODE_FONT)

    def _set_run_shading(self, run, fill: str):
        r_pr = run._element.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        r_pr.append(shd)

    def _set_cell_shading(self, cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _set_cell_margins(self, cell, *, top=100, start=160, bottom=100, end=160):
        """单元格内边距，单位 twips（约 8pt ≈ 160 twips）。"""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(val))
            node.set(qn("w:type"), "dxa")
            tc_mar.append(node)
        tc_pr.append(tc_mar)

    def _set_code_block_accent_border(self, cell):
        """左侧蓝色强调线，其余边框无。"""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "0")
        left.set(qn("w:color"), _CODE_BLOCK_ACCENT)
        tc_borders.append(left)
        for side in ("top", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "nil")
            tc_borders.append(border)
        tc_pr.append(tc_borders)

    def _set_table_no_outer_border(self, table):
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        tbl_borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "nil")
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def _add_inline_code_run(self, paragraph, text: str):
        run = paragraph.add_run(text)
        self._set_run_mono_font(run, inline=True)
        self._set_run_shading(run, _CODE_INLINE_BG)

    def _add_code_block(self, doc, code: str, _language: str = ""):
        """用单列表格包裹代码块：浅灰底 + 左侧蓝色强调线。"""
        code = code.replace("\r\n", "\n").replace("\r", "\n").rstrip("\n")
        if not code:
            return

        table = doc.add_table(rows=1, cols=1)
        self._set_table_no_outer_border(table)
        table.columns[0].width = Inches(6.2)

        cell = table.cell(0, 0)
        self._set_cell_shading(cell, _CODE_BLOCK_BG)
        self._set_code_block_accent_border(cell)
        self._set_cell_margins(cell, top=50, start=100, bottom=50, end=100)

        cell._element.clear_content()

        lines = code.split("\n")
        for i, line in enumerate(lines):
            line_p = cell.add_paragraph()
            line_run = line_p.add_run(line if line else " ")
            self._set_run_mono_font(line_run, size_pt=9)
            line_p.paragraph_format.space_before = Pt(0)
            line_p.paragraph_format.space_after = Pt(0)
            line_p.paragraph_format.line_spacing = 1.0
    
    def _set_table_borders(self, table, color: str = _TABLE_BORDER):
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        tbl_borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def _is_table_header_cell(self, row_idx: int, html_cell) -> bool:
        return row_idx == 0 or html_cell.name == "th"

    def _fill_table_cell(self, doc_cell, text: str, *, header: bool = False, alt_row: bool = False):
        doc_cell.text = text
        self._set_cell_margins(doc_cell, top=80, start=120, bottom=80, end=120)

        if header:
            self._set_cell_shading(doc_cell, _TABLE_HEADER_BG)
            for paragraph in doc_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = _TABLE_HEADER_TEXT
                    run.font.size = Pt(10.5)
                    r_pr = run._element.get_or_add_rPr()
                    _set_element_fonts(
                        r_pr, east_asia=self.heading_config.font_name, bold=True
                    )
        else:
            if alt_row:
                self._set_cell_shading(doc_cell, _TABLE_BODY_ALT_BG)
            for paragraph in doc_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    self._apply_body_run_font(run)

    def _add_html_table(self, doc, table):
        rows = table.find_all("tr")
        if not rows:
            return

        header_cells = rows[0].find_all(["th", "td"])
        col_count = len(header_cells)
        if col_count == 0:
            return

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(0)

        doc_table = doc.add_table(rows=len(rows), cols=col_count)
        doc_table.style = "Table Grid"
        doc_table.autofit = True
        self._set_table_borders(doc_table)

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for j, cell in enumerate(cells):
                if j >= col_count:
                    continue
                is_header = self._is_table_header_cell(i, cell)
                self._fill_table_cell(
                    doc_table.cell(i, j),
                    cell.get_text(),
                    header=is_header,
                    alt_row=not is_header and i % 2 == 0,
                )

        tail = doc.add_paragraph()
        tail.paragraph_format.space_before = Pt(6)
        tail.paragraph_format.space_after = Pt(0)
    
    def _download_image(self, url: str) -> io.BytesIO:
        """下载图片并返回字节流"""
        # 处理 URL 中的中文字符编码
        parsed = urlparse(url)
        # 对路径部分进行编码
        encoded_path = '/'.join(quote(segment, safe='/') for segment in parsed.path.split('/'))
        encoded_url = urlunparse((
            parsed.scheme,
            parsed.netloc,
            encoded_path,
            parsed.params,
            parsed.query,
            parsed.fragment
        ))
        
        # 配置请求头
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        # 禁用 SSL 警告（对于自签名证书）
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            # 下载图片，禁用 SSL 验证（适用于内网自签名证书）
            response = requests.get(
                encoded_url,
                headers=headers,
                timeout=30,
                verify=False  # 禁用 SSL 验证，适用于内网环境
            )
        
        response.raise_for_status()
        
        # 验证是否为图片内容
        content_type = response.headers.get('Content-Type', '')
        if not content_type.startswith('image/'):
            raise ValueError(f"URL 返回的不是图片类型: {content_type}")
        
        return io.BytesIO(response.content)
    
    def _add_image(self, doc, img_element, base_path: str | Path | None = None):
        """添加图片到Word文档"""
        src = img_element.get('src')
        alt = img_element.get('alt', '')
        
        if not src:
            return
        
        try:
            # 下载图片
            image_stream = self._load_image_stream(src, base_path=base_path)
            
            # 添加到Word文档
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_stream)
            
            # 如果有alt文本，添加为图片说明
            if alt:
                caption_paragraph = doc.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_body_run(caption_paragraph, alt, italic=True)
                
        except requests.exceptions.SSLError:
            error_paragraph = doc.add_paragraph()
            self._add_body_run(error_paragraph, f"[图片加载失败 - SSL错误: {alt or src}]", italic=True)
        except requests.exceptions.Timeout:
            error_paragraph = doc.add_paragraph()
            self._add_body_run(error_paragraph, f"[图片加载失败 - 请求超时: {alt or src}]", italic=True)
        except requests.exceptions.RequestException:
            error_paragraph = doc.add_paragraph()
            self._add_body_run(error_paragraph, f"[图片加载失败 - 网络错误: {alt or src}]", italic=True)
        except Exception as e:
            error_paragraph = doc.add_paragraph()
            self._add_body_run(
                error_paragraph, f"[图片加载失败: {alt or src} - {str(e)}]", italic=True
            )


class MarkdownToDocxTransformer:
    def __init__(self, heading_config: HeadingStyleConfig | None = None):
        self.heading_config = heading_config
        self.doc_tool = DocTool(heading_config=heading_config)

    def transform(
        self,
        markdown_content: str,          
        output_path: str,
        title: str = "Document",
        base_path: str | Path | None = None,
        heading_config: HeadingStyleConfig | None = None,
    ):
        params = {
            "markdown_content": markdown_content,
            "title": title,
            "base_path": str(base_path) if base_path is not None else None,
            "heading_config": heading_config or self.heading_config,
        }
        results = list(self.doc_tool.invoke(params))
        for result in results:
            if result["type"] == "text":
                pass
            elif result["type"] == "blob":
                with open(output_path, "wb") as f:
                    f.write(result["data"])

    def transform_from_source(
        self,
        source: str,
        output_path: str,
        *,
        title: str | None = None,
        encoding: str = "utf-8",
        timeout: float = 30.0,
    ) -> None:
        """从 Markdown 文件路径或 URL 读取内容并转换为 DOCX。

        :param source: 本地文件路径或 Markdown 的 HTTP/HTTPS URL
        :param output_path: 输出 DOCX 文件路径
        :param title: 文档标题；未传时从文件名或 URL 路径推导
        :param encoding: 本地文件读取编码，默认 utf-8
        :param timeout: URL 请求超时秒数
        """
        if source.startswith(("http://", "https://")):
            response = requests.get(source, timeout=timeout)
            response.raise_for_status()
            response.encoding = response.encoding or "utf-8"
            markdown_content = response.text
            source_base_path: str | Path | None = source.rsplit("/", 1)[0] if "/" in source else source
            if title is None:
                title = Path(urlparse(source).path).stem or "Document"
        else:
            path = Path(source)
            if not path.exists():
                raise FileNotFoundError(f"Not found file: {source}")
            markdown_content = path.read_text(encoding=encoding)
            source_base_path = path.parent
            if title is None:
                title = path.stem or "Document"
        self.transform(markdown_content, output_path, title, base_path=source_base_path)
    
    